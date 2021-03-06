import time,random,requests,multiprocessing,os
from bs4 import BeautifulSoup as bs
from docx import Document

allhead=['Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/67.0.3396.99 Safari/537.36',
         'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/39.0.2171.95 Safari/537.36 OPR/26.0.1656.60',
         'Opera/8.0 (Windows NT 5.1; U; en)',
        'Mozilla/5.0 (Windows NT 5.1; U; en; rv:1.8.1) Gecko/20061208 Firefox/2.0.0 Opera 9.50',
        'Mozilla/4.0 (compatible; MSIE 6.0; Windows NT 5.1; en) Opera 9.50',
        'Mozilla/5.0 (Windows NT 6.1; WOW64; rv:34.0) Gecko/20100101 Firefox/34.0',
        'Mozilla/5.0 (X11; U; Linux x86_64; zh-CN; rv:1.9.2.10) Gecko/20100922 Ubuntu/10.10 (maverick) Firefox/3.6.10',
         'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/30.0.1599.101 Safari/537.36',
        'Mozilla/5.0 (Windows NT 6.1; WOW64; Trident/7.0; rv:11.0) like Gecko'
         ]

def get_menu(url):      #获得目录以及链接
    menu=[]
    link=[]
    title=[]
    aa = requests.get(url, headers={'user-agent': random.choice(allhead)})
    bb = bs(aa.content, 'lxml')
    cc=bb.find_all(class_='box_con')
    dd=cc[1].find_all('dd')
    for i in dd:
        menu.append(i.text)     #获得目录文本
        link.append(url+i.a['href'])    #获得目录对应链接
    ee=bb.find(id='info')       #获得标题信息
    title=ee.text       #获得标题文本
    return menu,link,title

def get_content(i,lin,name,tar):    #建立编号.txt文档，下载内容
    try:
        time.sleep(random.uniform(0.6,1))
        aa1 = requests.get(lin, headers={'user-agent': random.choice(allhead)})
        time.sleep(random.uniform(0.6, 1))
        bb1 = bs(aa1.content, 'lxml')
        cc1 = bb1.find(id='content')
        dd1=bb1.find(class_='bookname')
        with open(tar+'/'+str(i)+'.txt','a') as f:
            f.write(dd1.find('h1').text+'\n')
            f.write(cc1.text.replace('\xa0',''))
            f.write( '\n')
            print('章节',i,'下载成功！')
    except:
        with open(tar+'/'+str(i)+'.txt','a') as f:
            f.write(name+'\n')
            f.write('此章节内容错误')
        print('章节',i,'内容错误！')


def multi_work(menu,link,tar):      #输入目录，连接，文件夹
    p = multiprocessing.Pool(processes=15)
    for i,lin in enumerate(link):
        p.apply_async(get_content, args=(i,lin,menu[i],tar,))
    p.close()
    p.join()

def make_docx(title,menu,tar):     #建立word文件
    f=Document()
    f.add_heading(title,0)
    for root,dirs,files in os.walk(tar):
        print(len(files),'个文件')
    for i in range(len(files)):
        with open(tar+'\\'+str(i)+'.txt') as n:
            f.add_heading(menu[i],1)
            f.add_paragraph(n.readlines())
    f.save(title.split('\n')[1]+'.docx')


if __name__=='__main__':
    url='http://www.biquge.jp/257312_41/'
    print('目录获取。。。。')
    menu,link,title=get_menu(url)     #获得目录以及链接
    print('目录获取成功。。。建立文件夹。。。')
    time.sleep(random.uniform(0.5,1))
    tar = os.getcwd()+'\\' + url.split('/')[-2]     #建立下载文件夹
    #os.makedirs(tar)
    print('开始下载。。。。')
    #multi_work(menu,link,tar)       #多线程下载
    print('下载结束。。。。制作word文档')
    make_docx(title,menu,tar)
    print('制作完成！')
